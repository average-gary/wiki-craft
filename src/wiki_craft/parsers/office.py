"""
Microsoft Office document parsers.

Handles Word (.docx) and Excel (.xlsx) files.
"""

import logging
from datetime import datetime
from pathlib import Path
from typing import BinaryIO, ClassVar

from wiki_craft.parsers.base import BaseParser
from wiki_craft.storage.models import (
    ContentBlock,
    ContentType,
    DocumentMetadata,
    DocumentType,
    ParsedDocument,
)

logger = logging.getLogger(__name__)


class WordParser(BaseParser):
    """Parser for Microsoft Word documents (.docx)."""

    supported_extensions: ClassVar[list[str]] = ["docx", "doc"]
    supported_mime_types: ClassVar[list[str]] = [
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        "application/msword",
    ]
    document_type: ClassVar[DocumentType] = DocumentType.WORD

    def parse(self, file_path: Path, file_content: BinaryIO | None = None) -> ParsedDocument:
        """
        Parse a Word document.

        Args:
            file_path: Path to the Word file
            file_content: Optional file-like object with content

        Returns:
            ParsedDocument with extracted content
        """
        from docx import Document
        from docx.opc.exceptions import PackageNotFoundError

        self.errors = []

        try:
            # Load document
            if file_content is not None:
                content_bytes = file_content.read()
                import io

                doc = Document(io.BytesIO(content_bytes))
                source_hash = self.compute_hash(content_bytes)
            else:
                doc = Document(file_path)
                source_hash = self.compute_file_hash(file_path)

            # Extract metadata
            metadata = self._extract_metadata(doc, file_path, source_hash)

            # Extract content
            content_blocks = []
            current_section: str | None = None
            section_hierarchy: list[str] = []
            position = 0

            for element in doc.element.body:
                # Handle paragraphs
                if element.tag.endswith("p"):
                    para = None
                    for p in doc.paragraphs:
                        if p._element is element:
                            para = p
                            break

                    if para is None or not para.text.strip():
                        continue

                    text = para.text.strip()
                    content_type = ContentType.PARAGRAPH

                    # Check if it's a heading
                    if para.style and para.style.name:
                        style_name = para.style.name.lower()
                        if "heading" in style_name or "title" in style_name:
                            content_type = ContentType.HEADING
                            current_section = text
                            # Update hierarchy based on heading level
                            level = self._get_heading_level(para.style.name)
                            section_hierarchy = section_hierarchy[: level - 1] + [text]

                    # Check for list items
                    if self._is_list_item(para):
                        content_type = ContentType.LIST

                    content_blocks.append(
                        ContentBlock(
                            text=text,
                            content_type=content_type,
                            section=current_section,
                            section_hierarchy=section_hierarchy.copy(),
                            position=position,
                        )
                    )
                    position += 1

                # Handle tables
                elif element.tag.endswith("tbl"):
                    for table in doc.tables:
                        if table._element is element:
                            table_text = self._extract_table(table)
                            if table_text:
                                content_blocks.append(
                                    ContentBlock(
                                        text=table_text,
                                        content_type=ContentType.TABLE,
                                        section=current_section,
                                        section_hierarchy=section_hierarchy.copy(),
                                        position=position,
                                    )
                                )
                                position += 1
                            break

            # Calculate word count
            metadata.word_count = sum(len(block.text.split()) for block in content_blocks)

            return ParsedDocument(
                metadata=metadata,
                content_blocks=content_blocks,
                raw_text="\n\n".join(block.text for block in content_blocks),
                parsing_errors=self.errors,
            )

        except PackageNotFoundError:
            self.add_error(f"Invalid or corrupted Word document: {file_path}")
            return ParsedDocument(
                metadata=DocumentMetadata(
                    source_path=str(file_path),
                    source_hash="",
                    filename=file_path.name,
                    document_type=DocumentType.WORD,
                ),
                parsing_errors=self.errors,
            )
        except Exception as e:
            self.add_error(f"Failed to parse Word document: {e}")
            raise

    def _extract_metadata(self, doc, file_path: Path, source_hash: str) -> DocumentMetadata:
        """Extract document metadata from Word file."""
        core_props = doc.core_properties

        return DocumentMetadata(
            source_path=str(file_path),
            source_hash=source_hash,
            filename=file_path.name,
            document_type=DocumentType.WORD,
            title=core_props.title or None,
            author=core_props.author or None,
            created_at=core_props.created,
            modified_at=core_props.modified,
        )

    def _get_heading_level(self, style_name: str) -> int:
        """Extract heading level from style name."""
        import re

        match = re.search(r"(\d+)", style_name)
        if match:
            return int(match.group(1))
        return 1

    def _is_list_item(self, para) -> bool:
        """Check if a paragraph is a list item."""
        # Check for numbering
        if para._element.pPr is not None:
            numPr = para._element.pPr.find(
                "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
            )
            if numPr is not None:
                return True
        # Check for bullet characters
        text = para.text.strip()
        if text and text[0] in "•-*◦▪":
            return True
        return False

    def _extract_table(self, table) -> str:
        """Extract table content as formatted text."""
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        return "\n".join(rows)


class ExcelParser(BaseParser):
    """Parser for Microsoft Excel spreadsheets (.xlsx)."""

    supported_extensions: ClassVar[list[str]] = ["xlsx", "xls"]
    supported_mime_types: ClassVar[list[str]] = [
        "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        "application/vnd.ms-excel",
    ]
    document_type: ClassVar[DocumentType] = DocumentType.EXCEL

    def parse(self, file_path: Path, file_content: BinaryIO | None = None) -> ParsedDocument:
        """
        Parse an Excel spreadsheet.

        Args:
            file_path: Path to the Excel file
            file_content: Optional file-like object with content

        Returns:
            ParsedDocument with extracted content
        """
        from openpyxl import load_workbook
        from openpyxl.utils.exceptions import InvalidFileException

        self.errors = []

        try:
            # Load workbook
            if file_content is not None:
                content_bytes = file_content.read()
                import io

                wb = load_workbook(io.BytesIO(content_bytes), data_only=True)
                source_hash = self.compute_hash(content_bytes)
            else:
                wb = load_workbook(file_path, data_only=True)
                source_hash = self.compute_file_hash(file_path)

            # Extract metadata
            metadata = self._extract_metadata(wb, file_path, source_hash)

            # Extract content from all sheets
            content_blocks = []
            position = 0

            for sheet_name in wb.sheetnames:
                sheet = wb[sheet_name]

                # Add sheet name as a heading
                content_blocks.append(
                    ContentBlock(
                        text=f"Sheet: {sheet_name}",
                        content_type=ContentType.HEADING,
                        section=sheet_name,
                        section_hierarchy=[sheet_name],
                        position=position,
                    )
                )
                position += 1

                # Extract table data
                table_text = self._extract_sheet_data(sheet)
                if table_text:
                    content_blocks.append(
                        ContentBlock(
                            text=table_text,
                            content_type=ContentType.TABLE,
                            section=sheet_name,
                            section_hierarchy=[sheet_name],
                            position=position,
                            metadata={"sheet_name": sheet_name},
                        )
                    )
                    position += 1

            wb.close()

            # Calculate word count
            metadata.word_count = sum(len(block.text.split()) for block in content_blocks)

            return ParsedDocument(
                metadata=metadata,
                content_blocks=content_blocks,
                raw_text="\n\n".join(block.text for block in content_blocks),
                parsing_errors=self.errors,
            )

        except InvalidFileException:
            self.add_error(f"Invalid or corrupted Excel file: {file_path}")
            return ParsedDocument(
                metadata=DocumentMetadata(
                    source_path=str(file_path),
                    source_hash="",
                    filename=file_path.name,
                    document_type=DocumentType.EXCEL,
                ),
                parsing_errors=self.errors,
            )
        except Exception as e:
            self.add_error(f"Failed to parse Excel file: {e}")
            raise

    def _extract_metadata(self, wb, file_path: Path, source_hash: str) -> DocumentMetadata:
        """Extract metadata from Excel workbook."""
        props = wb.properties

        created_at = None
        modified_at = None
        if props.created:
            created_at = (
                props.created if isinstance(props.created, datetime) else datetime.now()
            )
        if props.modified:
            modified_at = (
                props.modified if isinstance(props.modified, datetime) else datetime.now()
            )

        return DocumentMetadata(
            source_path=str(file_path),
            source_hash=source_hash,
            filename=file_path.name,
            document_type=DocumentType.EXCEL,
            title=props.title or None,
            author=props.creator or None,
            created_at=created_at,
            modified_at=modified_at,
        )

    def _extract_sheet_data(self, sheet) -> str:
        """Extract sheet data as formatted text."""
        rows = []
        for row in sheet.iter_rows(values_only=True):
            # Skip completely empty rows
            if all(cell is None for cell in row):
                continue
            cells = [str(cell) if cell is not None else "" for cell in row]
            rows.append(" | ".join(cells))
        return "\n".join(rows)
